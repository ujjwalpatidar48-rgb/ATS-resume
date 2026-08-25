import io
import json
import os
import re
import html
from pathlib import Path

import fitz
import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.platypus import SimpleDocTemplate, Paragraph, HRFlowable

# Optional AI support. The app still works in Smart Mode without an API key.
try:
    from openai import OpenAI
except Exception:
    OpenAI = None

st.set_page_config(
    page_title="ATS Resume Pro AI",
    page_icon="📄",
    layout="wide"
)

SKILLS = [
    "python", "java", "javascript", "typescript", "c", "c++", "c#",
    "sql", "mysql", "postgresql", "mongodb", "pandas", "numpy",
    "scikit-learn", "matplotlib", "seaborn", "statistics",
    "data analysis", "data analytics", "data visualization",
    "machine learning", "power bi", "tableau", "excel", "advanced excel",
    "dax", "html", "css", "react", "node.js", "django", "flask",
    "fastapi", "rest api", "api", "git", "github", "docker", "postman",
    "linux", "aws", "azure", "google cloud", "pytest", "unit testing",
    "etl", "data cleaning", "data preprocessing", "data pipeline",
    "oops", "object oriented programming", "json", "xml",
    "business analysis", "business analytics", "requirements gathering",
    "stakeholder management", "problem solving", "communication",
    "reporting", "dashboard", "dashboards", "data visualization",
    "ms excel", "power query", "pivot tables", "vlookup", "xlookup",
    "google sheets", "crm", "sap", "jira", "agile", "scrum"
]

SECTION_ALIASES = {
    "Professional Summary": ["summary", "professional summary", "profile", "objective", "career objective"],
    "Technical Skills": ["skills", "technical skills", "core skills", "technical expertise"],
    "Experience": ["experience", "work experience", "employment", "professional experience", "internship"],
    "Projects": ["projects", "project", "academic projects"],
    "Education": ["education", "academic", "qualification"],
    "Certifications": ["certification", "certifications", "courses"]
}


def normalize_text(text):
    text = (text or "").lower()
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def extract_skills(text):
    text = normalize_text(text)
    found = []
    for skill in SKILLS:
        if re.search(r"(?<!\w)" + re.escape(skill) + r"(?!\w)", text):
            found.append(skill)
    return sorted(set(found))


def extract_resume_text(uploaded_file):
    filename = uploaded_file.name.lower()
    if filename.endswith(".pdf"):
        pdf = fitz.open(stream=uploaded_file.read(), filetype="pdf")
        return "\n".join(page.get_text() for page in pdf)
    if filename.endswith(".docx"):
        doc = Document(uploaded_file)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if filename.endswith(".txt"):
        return uploaded_file.read().decode("utf-8", errors="ignore")
    raise ValueError("Only PDF, DOCX and TXT files are supported.")


def calculate_similarity(resume_text, jd_text):
    try:
        matrix = TfidfVectorizer(stop_words="english").fit_transform([resume_text, jd_text])
        return round(cosine_similarity(matrix[0:1], matrix[1:2])[0][0] * 100, 2)
    except Exception:
        return 0.0


def check_sections(text):
    text = normalize_text(text)
    return {
        section: any(alias in text for alias in aliases)
        for section, aliases in SECTION_ALIASES.items()
    }


def check_contact_information(text):
    return {
        "Email": bool(re.search(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b", text)),
        "Phone": bool(re.search(r"(\+91[\s-]?)?[6-9]\d{9}", text)),
        "LinkedIn": "linkedin.com" in text.lower(),
        "GitHub": "github.com" in text.lower()
    }


def formatting_score(text):
    score = 100
    suggestions = []
    if len(text.strip()) < 300:
        score -= 20
        suggestions.append("Resume content is too short.")
    if len(text) > 15000:
        score -= 10
        suggestions.append("Resume may be unnecessarily long.")
    if "@" not in text:
        score -= 10
        suggestions.append("Add a professional email address.")
    if not re.search(r"\d", text):
        score -= 5
        suggestions.append("Add measurable achievements where truthful.")
    if "\t\t" in text:
        score -= 5
        suggestions.append("Avoid complex multi-column layouts.")
    return max(score, 0), suggestions


def analyze_resume(resume_text, jd_text):
    resume_skills = set(extract_skills(resume_text))
    jd_skills = set(extract_skills(jd_text))
    matched = sorted(resume_skills & jd_skills)
    missing = sorted(jd_skills - resume_skills)

    keyword_score = round((len(matched) / len(jd_skills)) * 100, 2) if jd_skills else 0
    similarity_score = calculate_similarity(resume_text, jd_text)
    sections = check_sections(resume_text)
    section_score = round(sum(sections.values()) / len(sections) * 100, 2)
    contact = check_contact_information(resume_text)
    contact_score = round(sum(contact.values()) / len(contact) * 100, 2)
    format_score, format_suggestions = formatting_score(resume_text)

    ats_score = round(
        keyword_score * 0.45 +
        similarity_score * 0.25 +
        section_score * 0.15 +
        format_score * 0.10 +
        contact_score * 0.05,
        2
    )

    suggestions = []
    if missing:
        suggestions.append("Add only missing JD keywords that are genuinely supported by your background.")
    if not sections["Professional Summary"]:
        suggestions.append("Add a targeted Professional Summary.")
    if not sections["Technical Skills"]:
        suggestions.append("Create a dedicated Technical Skills section.")
    if not sections["Projects"]:
        suggestions.append("Add relevant projects.")
    if not sections["Experience"]:
        suggestions.append("Add internship, training or relevant experience if applicable.")
    suggestions.extend(format_suggestions)

    return {
        "ats_score": ats_score,
        "keyword_score": keyword_score,
        "similarity_score": similarity_score,
        "section_score": section_score,
        "format_score": format_score,
        "contact_score": contact_score,
        "matched_keywords": matched,
        "missing_keywords": missing,
        "jd_skills": sorted(jd_skills),
        "sections": sections,
        "contact": contact,
        "suggestions": suggestions
    }


def clean_bullets(value):
    lines = []
    for line in (value or "").splitlines():
        line = line.strip()
        if not line:
            continue
        line = re.sub(r"^[•●▪*-]\s*", "", line)
        lines.append("• " + line)
    return "\n".join(lines)


def extract_candidate_sections(text):
    """Best-effort section extraction used by Smart Mode and as AI grounding."""
    raw_lines = [x.strip() for x in (text or "").splitlines() if x.strip()]
    buckets = {k: [] for k in SECTION_ALIASES}
    current = None
    for line in raw_lines:
        low = normalize_text(line)
        matched_section = None
        for section, aliases in SECTION_ALIASES.items():
            if low in aliases or any(low == a for a in aliases):
                matched_section = section
                break
        if matched_section:
            current = matched_section
            continue
        if current:
            buckets[current].append(line)
    return {k: "\n".join(v) for k, v in buckets.items()}


def infer_role(jd_text):
    patterns = [
        r"(?:job title|role|position)\s*[:\-]\s*([^\n|,]+)",
        r"(?:hiring for|looking for|seeking)\s+(?:a|an)?\s*([A-Za-z][A-Za-z /&-]{2,50})"
    ]
    for pattern in patterns:
        m = re.search(pattern, jd_text, re.I)
        if m:
            return m.group(1).strip()
    common_roles = [
        "data analyst", "business analyst", "mis analyst", "bi analyst",
        "business intelligence analyst", "financial analyst", "software developer",
        "python developer", "full stack developer", "process analyst"
    ]
    low = normalize_text(jd_text)
    for role in common_roles:
        if role in low:
            return role.title()
    return "Target Role"


def smart_generate(resume_text, jd_text, analysis):
    """Deterministic fallback: never invents experience."""
    sections = extract_candidate_sections(resume_text)
    role = infer_role(jd_text)

    matched = analysis["matched_keywords"]
    existing_skills = extract_skills(resume_text)
    jd_only_supported = [s for s in matched if s in existing_skills]

    summary_seed = sections.get("Professional Summary", "")
    if summary_seed:
        summary = summary_seed.replace("\n", " ")
    else:
        skill_text = ", ".join([s.title() for s in jd_only_supported[:6]]) or ", ".join(
            s.title() for s in existing_skills[:6]
        )
        summary = (
            f"Detail-oriented professional targeting {role} opportunities with hands-on knowledge "
            f"in {skill_text}. Skilled in analyzing information, preparing reports, and using "
            f"data-driven approaches to support business decisions. Seeking to apply existing "
            f"skills in a role aligned with the requirements of the job description."
        )

    skills = ", ".join(s.title() for s in (jd_only_supported + [s for s in existing_skills if s not in jd_only_supported]))
    skills = skills[:1200]

    experience = sections.get("Experience", "")
    projects = sections.get("Projects", "")
    education = sections.get("Education", "")
    certifications = sections.get("Certifications", "")

    # Keep original factual content. Only reorder and emphasize supported JD skills.
    if experience:
        experience = clean_bullets(experience)
    if projects:
        projects = clean_bullets(projects)

    return {
        "name": "",
        "contact": "",
        "summary": summary,
        "skills": skills,
        "experience": experience,
        "projects": projects,
        "education": education,
        "certifications": certifications,
        "generation_note": "Smart Mode used: existing facts were preserved; no unsupported experience or skills were added."
    }


def get_secret(name):
    try:
        return st.secrets.get(name)
    except Exception:
        return None


def ai_generate(resume_text, jd_text, analysis):
    api_key = get_secret("OPENAI_API_KEY") or os.getenv("OPENAI_API_KEY")
    if not api_key or OpenAI is None:
        return None, "AI Mode is unavailable because OPENAI_API_KEY is not configured."

    model = get_secret("OPENAI_MODEL") or os.getenv("OPENAI_MODEL") or "gpt-4.1-mini"
    client = OpenAI(api_key=api_key)

    prompt = f"""
You are an expert ATS resume writer and recruiter.

TASK:
Create a one-page, ATS-friendly resume tailored to the supplied Job Description.

STRICT TRUTH RULES:
1. Use ONLY facts, skills, experience, projects, education and certifications present in the candidate resume.
2. Do NOT invent employers, dates, degrees, tools, achievements, percentages, responsibilities or certifications.
3. Do NOT add a JD keyword just because it appears in the JD. Add a skill/keyword only when the candidate resume supports it.
4. Remove irrelevant skills from the generated Technical Skills section when they are not useful for the target JD.
5. Rewrite the Professional Summary to match the target role using truthful candidate information.
6. Improve wording and ATS keyword alignment without changing facts.
7. Keep the output concise and suitable for one page.
8. Use standard ATS sections only: Professional Summary, Technical Skills, Experience, Projects, Education, Certifications.
9. Do not use tables, columns, icons or graphics.

Return ONLY valid JSON with these keys:
name, contact, summary, skills, experience, projects, education, certifications

CANDIDATE RESUME:
{resume_text}

JOB DESCRIPTION:
{jd_text}

CURRENT ATS ANALYSIS:
Matched keywords: {analysis["matched_keywords"]}
Missing JD keywords: {analysis["missing_keywords"]}
ATS score: {analysis["ats_score"]}
"""

    try:
        response = client.responses.create(
            model=model,
            input=prompt
        )
        content = response.output_text.strip()
        data = json.loads(content)
        required = ["name", "contact", "summary", "skills", "experience", "projects", "education", "certifications"]
        if not all(k in data for k in required):
            return None, "AI returned an incomplete resume."
        data["generation_note"] = "AI Mode used: resume wording and ordering were tailored to the JD while preserving candidate facts."
        return data, None
    except Exception as e:
        return None, f"AI generation failed: {e}"


def add_heading(document, text):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10.5)

    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_body(document, text):
    for line in (text or "").splitlines():
        line = line.strip()
        if not line:
            continue
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(1.5)
        if line.startswith(("•", "-", "*")):
            p.style = document.styles["List Bullet"]
            line = line.lstrip("•-* ")
        run = p.add_run(line)
        run.font.name = "Arial"
        run.font.size = Pt(9.2)


def build_docx(data):
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    document.styles["Normal"].font.name = "Arial"
    document.styles["Normal"].font.size = Pt(9.2)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(data.get("name", ""))
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(16)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(data.get("contact", ""))
    r.font.name = "Arial"
    r.font.size = Pt(8.5)

    for title, key in [
        ("Professional Summary", "summary"),
        ("Technical Skills", "skills"),
        ("Experience", "experience"),
        ("Projects", "projects"),
        ("Education", "education"),
        ("Certifications", "certifications")
    ]:
        value = data.get(key, "")
        if value:
            add_heading(document, title)
            add_body(document, value)

    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def build_pdf(data):
    stream = io.BytesIO()
    document = SimpleDocTemplate(
        stream, pagesize=A4,
        rightMargin=12 * mm, leftMargin=12 * mm,
        topMargin=10 * mm, bottomMargin=10 * mm
    )
    styles = getSampleStyleSheet()
    name_style = ParagraphStyle(
        "ResumeName", parent=styles["Normal"], fontName="Helvetica-Bold",
        fontSize=15, alignment=TA_CENTER, spaceAfter=3
    )
    contact_style = ParagraphStyle(
        "ResumeContact", parent=styles["Normal"], fontName="Helvetica",
        fontSize=8.5, alignment=TA_CENTER, spaceAfter=5
    )
    heading_style = ParagraphStyle(
        "ResumeHeading", parent=styles["Normal"], fontName="Helvetica-Bold",
        fontSize=9.5, spaceBefore=5, spaceAfter=2
    )
    body_style = ParagraphStyle(
        "ResumeBody", parent=styles["Normal"], fontName="Helvetica",
        fontSize=8.5, leading=10.5, spaceAfter=1.5
    )

    story = [
        Paragraph(html.escape(data.get("name", "")), name_style),
        Paragraph(html.escape(data.get("contact", "")), contact_style)
    ]

    for title, key in [
        ("PROFESSIONAL SUMMARY", "summary"),
        ("TECHNICAL SKILLS", "skills"),
        ("EXPERIENCE", "experience"),
        ("PROJECTS", "projects"),
        ("EDUCATION", "education"),
        ("CERTIFICATIONS", "certifications")
    ]:
        value = data.get(key, "")
        if not value:
            continue
        story.append(Paragraph(title, heading_style))
        story.append(HRFlowable(width="100%", thickness=0.5, spaceAfter=2))
        for line in value.splitlines():
            line = line.strip()
            if not line:
                continue
            if line.startswith(("•", "-", "*")):
                line = "• " + line.lstrip("•-* ")
            story.append(Paragraph(html.escape(line), body_style))

    document.build(story)
    return stream.getvalue()


# UI
st.title("📄 ATS Resume Pro AI")
st.caption("ATS Score + JD Analysis + Automatic JD-Tailored Resume Builder")

with st.sidebar:
    st.header("⚙️ Features")
    st.write("✅ ATS Score")
    st.write("✅ JD Keyword Matching")
    st.write("✅ Missing Keyword Analysis")
    st.write("✅ Automatic Resume Tailoring")
    st.write("✅ ATS-Friendly DOCX")
    st.write("✅ ATS-Friendly PDF")
    st.divider()
    st.info(
        "AI Mode never invents experience or skills. It uses only information supported by the uploaded resume. "
        "Smart Mode works without an API key."
    )

left, right = st.columns(2)

with left:
    st.header("1️⃣ Your Current Resume")
    uploaded_file = st.file_uploader("Upload Resume", type=["pdf", "docx", "txt"])
    uploaded_text = ""
    if uploaded_file:
        try:
            uploaded_text = extract_resume_text(uploaded_file)
            st.success("Resume successfully loaded.")
        except Exception as e:
            st.error(f"Error reading resume: {e}")

    resume_text = st.text_area(
        "Resume Text",
        value=uploaded_text,
        height=350,
        placeholder="Upload your resume or paste resume text here..."
    )

with right:
    st.header("2️⃣ Target Job Description")
    jd_text = st.text_area(
        "Paste Complete Job Description",
        height=350,
        placeholder="Paste the complete job description here..."
    )

if st.button("🔍 Analyze ATS Score", type="primary", use_container_width=True):
    if not resume_text.strip():
        st.warning("Please upload or paste your resume.")
    elif not jd_text.strip():
        st.warning("Please paste the Job Description.")
    else:
        st.session_state["analysis"] = analyze_resume(resume_text, jd_text)

if "analysis" in st.session_state:
    result = st.session_state["analysis"]
    st.divider()
    st.header("3️⃣ ATS Analysis")

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("ATS Score", f'{result["ats_score"]}%')
    c2.metric("Keyword Match", f'{result["keyword_score"]}%')
    c3.metric("JD Similarity", f'{result["similarity_score"]}%')
    c4.metric("Format Score", f'{result["format_score"]}%')

    if result["ats_score"] >= 85:
        st.success("🔥 Excellent ATS alignment.")
    elif result["ats_score"] >= 70:
        st.success("👍 Good ATS alignment. Tailoring can improve it further.")
    elif result["ats_score"] >= 50:
        st.warning("⚠️ Moderate ATS alignment. Tailor the resume more closely.")
    else:
        st.error("❌ Low ATS alignment. Significant tailoring is recommended.")

    col1, col2 = st.columns(2)
    with col1:
        st.subheader("✅ Matched Keywords")
        st.write(", ".join(result["matched_keywords"]) if result["matched_keywords"] else "No tracked keywords matched.")
    with col2:
        st.subheader("⚠️ Missing JD Keywords")
        if result["missing_keywords"]:
            for kw in result["missing_keywords"]:
                st.write(f"• {kw}")
        else:
            st.success("No major tracked keywords are missing.")

    st.subheader("💡 Improvement Suggestions")
    for item in result["suggestions"]:
        st.write(f"• {item}")

    st.divider()
    st.header("4️⃣ Automatically Create JD-Tailored Resume")

    mode = st.radio(
        "Generation Mode",
        ["Smart Mode (free, no API key)", "AI Mode (best tailoring, requires OpenAI API key)"],
        horizontal=True
    )

    st.warning(
        "Important: The generator will not add unsupported skills, experience, achievements or certifications. "
        "Only truthful information from your resume can be used."
    )

    if st.button("✨ Generate ATS-Friendly Resume for This JD", type="primary", use_container_width=True):
        if mode.startswith("AI"):
            generated, error = ai_generate(resume_text, jd_text, result)
            if generated is None:
                st.warning(error)
                st.info("Using Smart Mode instead so you can still generate a resume.")
                generated = smart_generate(resume_text, jd_text, result)
        else:
            generated = smart_generate(resume_text, jd_text, result)

        st.session_state["generated_resume"] = generated
        st.success("🎉 JD-tailored resume generated successfully.")

if "generated_resume" in st.session_state:
    data = st.session_state["generated_resume"]
    st.divider()
    st.header("5️⃣ Generated Resume")

    st.caption(data.get("generation_note", ""))

    st.text_input("Name", value=data.get("name", ""), key="out_name")
    st.text_input("Contact", value=data.get("contact", ""), key="out_contact")
    st.text_area("Professional Summary", value=data.get("summary", ""), height=130, key="out_summary")
    st.text_area("Technical Skills", value=data.get("skills", ""), height=100, key="out_skills")
    st.text_area("Experience / Internship", value=data.get("experience", ""), height=170, key="out_experience")
    st.text_area("Projects", value=data.get("projects", ""), height=170, key="out_projects")
    st.text_area("Education", value=data.get("education", ""), height=110, key="out_education")
    st.text_area("Certifications", value=data.get("certifications", ""), height=110, key="out_certifications")

    final_data = {
        "name": st.session_state.get("out_name", data.get("name", "")),
        "contact": st.session_state.get("out_contact", data.get("contact", "")),
        "summary": st.session_state.get("out_summary", data.get("summary", "")),
        "skills": st.session_state.get("out_skills", data.get("skills", "")),
        "experience": st.session_state.get("out_experience", data.get("experience", "")),
        "projects": st.session_state.get("out_projects", data.get("projects", "")),
        "education": st.session_state.get("out_education", data.get("education", "")),
        "certifications": st.session_state.get("out_certifications", data.get("certifications", ""))
    }

    docx_file = build_docx(final_data)
    pdf_file = build_pdf(final_data)

    c1, c2 = st.columns(2)
    with c1:
        st.download_button(
            "⬇️ Download ATS Resume DOCX",
            data=docx_file,
            file_name="JD_Tailored_ATS_Resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    with c2:
        st.download_button(
            "⬇️ Download ATS Resume PDF",
            data=pdf_file,
            file_name="JD_Tailored_ATS_Resume.pdf",
            mime="application/pdf",
            use_container_width=True
        )

st.divider()
st.caption("ATS Resume Pro AI | Tailor honestly. Never add skills or achievements you do not genuinely have.")
